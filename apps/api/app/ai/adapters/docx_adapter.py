"""
DocuMind AI — DOCX Adapter (python-docx).

Handles .docx files using python-docx — lightweight, no ML dependencies,
~20 MB install, suitable for Render's 512 MB limit.
"""

from __future__ import annotations

import logging
import uuid

from app.ai.adapters.base import BaseAdapter
from app.ai.models import BlockType, DocumentBlock, NormalizedDocument

logger = logging.getLogger(__name__)


class DocxAdapter(BaseAdapter):
    """Lightweight DOCX adapter using python-docx. No PyTorch, no ML models."""

    def parse(self, file_bytes: bytes) -> NormalizedDocument:
        try:
            import io

            import docx  # python-docx
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is not installed. Add 'python-docx>=1.1.0' to requirements.txt."
            ) from exc

        logger.info(f"[PROCESSING] docx_adapter.parse started document_id={self.document_id}")

        blocks: list[DocumentBlock] = []

        try:
            doc = docx.Document(io.BytesIO(file_bytes))

            # Estimate paragraph count as pseudo-page-count (1 page ~ 40 paragraphs heuristic)
            para_count = len(doc.paragraphs)
            page_count = max(1, para_count // 40)

            for para in doc.paragraphs:
                content = para.text.strip()
                if not content:
                    continue

                # Use paragraph style to classify blocks
                style_name = (para.style.name or "").lower()
                if "heading" in style_name:
                    block_type = BlockType.HEADING
                else:
                    block_type = BlockType.TEXT

                blocks.append(DocumentBlock(
                    block_id=str(uuid.uuid4()),
                    block_type=block_type,
                    content=content,
                    page_number=None,  # python-docx does not expose page breaks easily
                    bbox=None,
                ))

            # Extract tables as text blocks
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                table_text = "\n".join(rows).strip()
                if table_text:
                    blocks.append(DocumentBlock(
                        block_id=str(uuid.uuid4()),
                        block_type=BlockType.TABLE,
                        content=table_text,
                        page_number=None,
                        bbox=None,
                    ))

        except Exception as exc:
            logger.error(
                f"[PROCESSING] docx_adapter.parse failed"
                f" document_id={self.document_id} error={exc}"
            )
            raise RuntimeError(f"DOCX parsing failed: {exc}") from exc

        logger.info(
            f"[PROCESSING] docx_adapter.parse completed document_id={self.document_id} "
            f"est_pages={page_count} blocks={len(blocks)}"
        )

        return NormalizedDocument(
            document_id=self.document_id,
            user_id=self.user_id,
            file_path=self.file_path,
            source_mime_type=self.mime_type,
            page_count=page_count,
            title=self.title,
            blocks=blocks,
            processing_metadata={"adapter": "DocxAdapter"},
        )
